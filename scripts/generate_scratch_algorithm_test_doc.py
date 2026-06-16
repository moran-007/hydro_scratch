from __future__ import annotations

import json
from dataclasses import dataclass
from pathlib import Path
from typing import Callable

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "docs"
DOC_DATE = "2026-06-16"
DOC_STEM = "scratch-algorithm-ask-answer-test-20260616"


@dataclass(frozen=True)
class Case:
    name: str
    answers: list[str]
    expected: str
    score: int
    hidden: bool = False
    note: str = ""


@dataclass(frozen=True)
class Problem:
    pid: str
    title: str
    level: str
    goal: str
    ask_flow: list[str]
    output_rule: str
    compare_mode: str
    cases: list[Case]
    solver: Callable[[list[str]], str]
    scratch_tip: str


def solve_sum_two(values: list[str]) -> str:
    return str(int(values[0]) + int(values[1]))


def solve_max_three(values: list[str]) -> str:
    return str(max(int(item) for item in values[:3]))


def solve_even_odd(values: list[str]) -> str:
    return "even" if int(values[0]) % 2 == 0 else "odd"


def solve_sum_to_n(values: list[str]) -> str:
    n = int(values[0])
    return str(n * (n + 1) // 2)


def solve_pass_count(values: list[str]) -> str:
    n = int(values[0])
    scores = [int(item) for item in values[1 : 1 + n]]
    return str(sum(1 for item in scores if item >= 60))


def solve_word_length(values: list[str]) -> str:
    return str(len(values[0]))


def solve_c_to_f(values: list[str]) -> str:
    c = float(values[0])
    f = c * 9 / 5 + 32
    if f.is_integer():
        return str(int(f))
    return f"{f:.6f}".rstrip("0").rstrip(".")


def solve_min_max(values: list[str]) -> str:
    n = int(values[0])
    numbers = [int(item) for item in values[1 : 1 + n]]
    return f"{min(numbers)} {max(numbers)}"


PROBLEMS: list[Problem] = [
    Problem(
        pid="SA01",
        title="两数求和",
        level="入门",
        goal="输入两个整数，输出它们的和。",
        ask_flow=["询问第一个整数 a", "询问第二个整数 b"],
        output_rule="输出 a + b 的结果。",
        compare_mode="number",
        scratch_tip="每次询问后立即把“答案”保存到变量 a 或 b，最后把 a+b 设置到 output。",
        solver=solve_sum_two,
        cases=[
            Case("公开 1：普通正数", ["1", "2"], "3", 20),
            Case("公开 2：零值", ["0", "0"], "0", 20),
            Case("公开 3：一正一负", ["-5", "8"], "3", 20),
            Case("隐藏 1：进位", ["999", "1"], "1000", 20, True),
            Case("隐藏 2：两个负数", ["-100", "-250"], "-350", 20, True),
        ],
    ),
    Problem(
        pid="SA02",
        title="三个数最大值",
        level="入门",
        goal="输入三个整数，输出其中最大的一个。",
        ask_flow=["询问第一个整数 a", "询问第二个整数 b", "询问第三个整数 c"],
        output_rule="输出 max(a, b, c)。",
        compare_mode="number",
        scratch_tip="可以先把 max 设为 a，再依次和 b、c 比较并更新。",
        solver=solve_max_three,
        cases=[
            Case("公开 1：中间最大", ["3", "9", "5"], "9", 20),
            Case("公开 2：全为负数", ["-1", "-5", "-3"], "-1", 20),
            Case("公开 3：存在相等", ["8", "8", "7"], "8", 20),
            Case("隐藏 1：后两项相等最大", ["0", "100", "100"], "100", 20, True),
            Case("隐藏 2：第二项最大", ["-10", "20", "15"], "20", 20, True),
        ],
    ),
    Problem(
        pid="SA03",
        title="奇偶判断",
        level="入门",
        goal="输入一个整数，判断它是偶数还是奇数。",
        ask_flow=["询问一个整数 n"],
        output_rule="如果 n 是偶数，输出 even；否则输出 odd。",
        compare_mode="trim",
        scratch_tip="使用“n 除以 2 的余数”判断；输出必须是小写英文 even 或 odd。",
        solver=solve_even_odd,
        cases=[
            Case("公开 1：零", ["0"], "even", 20),
            Case("公开 2：正奇数", ["7"], "odd", 20),
            Case("公开 3：负偶数", ["-4"], "even", 20),
            Case("隐藏 1：较大奇数", ["101"], "odd", 20, True),
            Case("隐藏 2：较大偶数", ["1000"], "even", 20, True),
        ],
    ),
    Problem(
        pid="SA04",
        title="从 1 加到 n",
        level="循环",
        goal="输入整数 n，输出 1+2+...+n 的和。",
        ask_flow=["询问整数 n"],
        output_rule="输出从 1 累加到 n 的总和；n=0 时输出 0。",
        compare_mode="number",
        scratch_tip="使用重复执行或重复执行直到，配合 i 和 sum 两个变量完成累加。",
        solver=solve_sum_to_n,
        cases=[
            Case("公开 1：最小正数", ["1"], "1", 20),
            Case("公开 2：小范围循环", ["5"], "15", 20),
            Case("公开 3：十项求和", ["10"], "55", 20),
            Case("隐藏 1：零", ["0"], "0", 20, True),
            Case("隐藏 2：一百项", ["100"], "5050", 20, True),
        ],
    ),
    Problem(
        pid="SA05",
        title="统计及格人数",
        level="循环",
        goal="先输入人数 n，再输入 n 个成绩，输出大于等于 60 分的人数。",
        ask_flow=["询问人数 n", "循环 n 次，依次询问每个学生成绩"],
        output_rule="输出成绩 >= 60 的数量。",
        compare_mode="number",
        scratch_tip="循环读取每个成绩，若成绩 >= 60，则 passCount 增加 1。",
        solver=solve_pass_count,
        cases=[
            Case("公开 1：边界 60", ["3", "60", "59", "100"], "2", 20),
            Case("公开 2：混合成绩", ["5", "0", "100", "60", "61", "59"], "3", 20),
            Case("公开 3：单人不及格", ["1", "59"], "0", 20),
            Case("隐藏 1：全部及格", ["4", "60", "60", "60", "60"], "4", 20, True),
            Case("隐藏 2：全部不及格", ["6", "10", "20", "30", "40", "50", "59"], "0", 20, True),
        ],
    ),
    Problem(
        pid="SA06",
        title="单词长度",
        level="字符串",
        goal="输入一个不含空格的单词，输出它的字符个数。",
        ask_flow=["询问一个不含空格的单词 word"],
        output_rule="输出 word 的长度。",
        compare_mode="number",
        scratch_tip="使用运算类积木“word 的长度”，不要把中文提示文字也拼进 output。",
        solver=solve_word_length,
        cases=[
            Case("公开 1：短单词", ["cat"], "3", 20),
            Case("公开 2：Scratch", ["Scratch"], "7", 20),
            Case("公开 3：单字符", ["A"], "1", 20),
            Case("隐藏 1：字母数字混合", ["hello2026"], "9", 20, True),
            Case("隐藏 2：长单词", ["algorithm"], "9", 20, True),
        ],
    ),
    Problem(
        pid="SA07",
        title="摄氏温度转华氏温度",
        level="表达式",
        goal="输入摄氏温度 c，输出对应华氏温度 f。",
        ask_flow=["询问摄氏温度 c"],
        output_rule="按公式 f = c * 9 / 5 + 32 输出结果。",
        compare_mode="number",
        scratch_tip="可以直接使用 Scratch 数学表达式；小数结果允许按数字比较。",
        solver=solve_c_to_f,
        cases=[
            Case("公开 1：冰点", ["0"], "32", 20),
            Case("公开 2：沸点", ["100"], "212", 20),
            Case("公开 3：相等点", ["-40"], "-40", 20),
            Case("隐藏 1：体温附近", ["37"], "98.6", 20, True),
            Case("隐藏 2：常温", ["25"], "77", 20, True),
        ],
    ),
    Problem(
        pid="SA08",
        title="一组数的最小值和最大值",
        level="循环与比较",
        goal="先输入 n，再输入 n 个整数，输出最小值和最大值。",
        ask_flow=["询问整数个数 n", "循环 n 次，依次询问每个整数"],
        output_rule="输出格式为“最小值 最大值”，中间用一个空格分隔。",
        compare_mode="tokens",
        scratch_tip="读第一个数时同时初始化 min 和 max，后续每读一个数就分别比较更新。",
        solver=solve_min_max,
        cases=[
            Case("公开 1：普通序列", ["5", "3", "1", "9", "-2", "8"], "-2 9", 20),
            Case("公开 2：单个数", ["1", "7"], "7 7", 20),
            Case("公开 3：全为负数", ["4", "-5", "-1", "-9", "-3"], "-9 -1", 20),
            Case("隐藏 1：存在重复", ["6", "2", "2", "2", "9", "9", "0"], "0 9", 20, True),
            Case("隐藏 2：范围较大", ["5", "100", "-100", "50", "0", "99"], "-100 100", 20, True),
        ],
    ),
]


def validate_cases() -> list[dict[str, str | int | bool]]:
    results: list[dict[str, str | int | bool]] = []
    for problem in PROBLEMS:
        for case in problem.cases:
            actual = problem.solver(case.answers)
            passed = actual == case.expected
            results.append(
                {
                    "pid": problem.pid,
                    "title": problem.title,
                    "case": case.name,
                    "input": "\n".join(case.answers),
                    "expected": case.expected,
                    "actual": actual,
                    "score": case.score,
                    "hidden": case.hidden,
                    "passed": passed,
                }
            )
            if not passed:
                raise AssertionError(
                    f"{problem.pid} {case.name}: expected {case.expected!r}, got {actual!r}"
                )
    return results


def quick_line(case: Case) -> str:
    prefix = "* " if case.hidden else ""
    input_text = "\\n".join(case.answers)
    return f"{prefix}{input_text} => {case.expected} => {case.score} => {case.name}"


def judge_config(problem: Problem) -> dict:
    return {
        "schemaVersion": 2,
        "totalScore": 100,
        "algorithm": {
            "target": "Stage",
            "inputVariable": "input",
            "outputVariable": "output",
            "compareMode": problem.compare_mode,
            "numericTolerance": 1e-6,
            "waitMs": 1000,
            "timeoutMs": 6000,
            "cases": [
                {
                    "name": case.name,
                    "input": "\n".join(case.answers),
                    "expectedOutput": case.expected,
                    "score": case.score,
                    **({"hidden": True} if case.hidden else {}),
                }
                for case in problem.cases
            ],
        },
    }


def write_json() -> Path:
    path = OUT_DIR / f"{DOC_STEM}.json"
    payload = {
        "generatedAt": DOC_DATE,
        "site": "https://moran007.top/d/scratch/",
        "note": "Passwords are intentionally not stored in this repository document.",
        "problems": [
            {
                "pid": problem.pid,
                "title": problem.title,
                "level": problem.level,
                "goal": problem.goal,
                "askFlow": problem.ask_flow,
                "outputRule": problem.output_rule,
                "compareMode": problem.compare_mode,
                "quickEntry": "\n".join(quick_line(case) for case in problem.cases),
                "judgeConfig": judge_config(problem),
            }
            for problem in PROBLEMS
        ],
    }
    path.write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8")
    return path


def write_quick_entries() -> Path:
    path = OUT_DIR / f"{DOC_STEM}-quick-entry.txt"
    blocks: list[str] = []
    for problem in PROBLEMS:
        blocks.append(f"{problem.pid} {problem.title} / compareMode={problem.compare_mode}")
        blocks.extend(quick_line(case) for case in problem.cases)
        blocks.append("")
    path.write_text("\n".join(blocks), encoding="utf-8")
    return path


def markdown_table(rows: list[list[str]]) -> str:
    if not rows:
        return ""
    header = rows[0]
    body = rows[1:]
    lines = [
        "| " + " | ".join(header) + " |",
        "| " + " | ".join("---" for _ in header) + " |",
    ]
    for row in body:
        lines.append("| " + " | ".join(item.replace("\n", "<br>") for item in row) + " |")
    return "\n".join(lines)


def write_markdown(results: list[dict[str, str | int | bool]]) -> Path:
    path = OUT_DIR / f"{DOC_STEM}.md"
    lines: list[str] = [
        "# Scratch 算法题询问/回答测试文档",
        "",
        f"生成日期：{DOC_DATE}",
        "",
        "目标站点：https://moran007.top/d/scratch/",
        "",
        "账号记录：管理员用户名 yang；学生用户名 ceshi1。密码不写入本地测试文档，请使用本次测试提供的口令。",
        "",
        "## 测试约定",
        "",
        "本组题以 Scratch 的“询问并等待 / 答案”作为人工测试输入方式。每个测试点的“回答序列”表示测试人员应按顺序输入的答案。",
        "",
        "若使用 Hydro Scratch 插件的算法题自动评测，可把同一回答序列用换行符拼成变量 input，学生程序最终把纯结果写入变量 output。文档中同时给出可粘贴到“算法题测试点快速录入”的格式。",
        "",
        "输出应只包含答案本身，不要带“答案是：”等提示文字。",
        "",
        "## 本地参考运行结果",
        "",
    ]
    summary_rows = [["题号", "题目", "用例数", "通过", "总分", "结果"]]
    for problem in PROBLEMS:
        problem_results = [row for row in results if row["pid"] == problem.pid]
        passed = sum(1 for row in problem_results if row["passed"])
        total_score = sum(int(row["score"]) for row in problem_results)
        summary_rows.append(
            [
                problem.pid,
                problem.title,
                str(len(problem_results)),
                f"{passed}/{len(problem_results)}",
                str(total_score),
                "PASS" if passed == len(problem_results) else "FAIL",
            ]
        )
    lines.append(markdown_table(summary_rows))
    lines.append("")
    lines.append(f"本地参考求解器共校验 {len(results)} 个测试点，全部通过。")
    lines.append("")

    for problem in PROBLEMS:
        lines.extend(
            [
                f"## {problem.pid} {problem.title}",
                "",
                f"难度：{problem.level}",
                "",
                f"题意：{problem.goal}",
                "",
                f"询问顺序：{'；'.join(problem.ask_flow)}。",
                "",
                f"输出要求：{problem.output_rule}",
                "",
                f"Scratch 实现提示：{problem.scratch_tip}",
                "",
                f"比较方式：{problem.compare_mode}",
                "",
            ]
        )
        case_rows = [["用例", "回答序列", "期望输出", "分值", "结果", "说明"]]
        for case in problem.cases:
            actual = problem.solver(case.answers)
            case_rows.append(
                [
                    case.name,
                    "<br>".join(case.answers),
                    case.expected,
                    str(case.score),
                    "PASS" if actual == case.expected else "FAIL",
                    "隐藏" if case.hidden else "公开",
                ]
            )
        lines.append(markdown_table(case_rows))
        lines.extend(["", "快速录入：", "", "```text"])
        lines.extend(quick_line(case) for case in problem.cases)
        lines.extend(["```", ""])
    path.write_text("\n".join(lines), encoding="utf-8")
    return path


def set_run_font(run, size: float | None = None, bold: bool | None = None, color: str | None = None):
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.start_type = WD_SECTION.NEW_PAGE
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    heading_specs = {
        "Heading 1": (16, "2E74B5", 18, 10),
        "Heading 2": (13, "2E74B5", 14, 7),
        "Heading 3": (12, "1F4D78", 10, 5),
    }
    for style_name, (size, color, before, after) in heading_specs.items():
        style = styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)

    header = section.header.paragraphs[0]
    header.text = "Scratch 算法题询问/回答测试文档"
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for run in header.runs:
        set_run_font(run, 9, False, "555555")

    footer = section.footer.paragraphs[0]
    footer.text = f"Generated {DOC_DATE}"
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for run in footer.runs:
        set_run_font(run, 9, False, "555555")


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin_name, value in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = tc_mar.find(qn(f"w:{margin_name}"))
        if node is None:
            node = OxmlElement(f"w:{margin_name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths: list[int], indent: int = 120) -> None:
    table.autofit = False
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent))
    tbl_ind.set(qn("w:type"), "dxa")

    existing_grid = tbl.tblGrid
    if existing_grid is not None:
        tbl.remove(existing_grid)
    grid = OxmlElement("w:tblGrid")
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    tbl.insert(1, grid)

    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            cell.width = Inches(width / 1440)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_margins(cell)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")


def set_cell_text(cell, text: str, size: float = 9.5, bold: bool = False, color: str | None = None):
    cell.text = ""
    parts = text.split("\n")
    for index, part in enumerate(parts):
        paragraph = cell.paragraphs[0] if index == 0 else cell.add_paragraph()
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.line_spacing = 1.15
        run = paragraph.add_run(part)
        set_run_font(run, size, bold, color)


def add_table(doc: Document, rows: list[list[str]], widths: list[int], header_fill: str = "E8EEF5"):
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = "Table Grid"
    for r_idx, row in enumerate(rows):
        for c_idx, text in enumerate(row):
            cell = table.rows[r_idx].cells[c_idx]
            set_cell_text(cell, text, bold=(r_idx == 0))
            if r_idx == 0:
                shade_cell(cell, header_fill)
    set_table_geometry(table, widths)
    return table


def add_code_block(doc: Document, text: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell = table.cell(0, 0)
    shade_cell(cell, "F4F6F9")
    set_cell_margins(cell, top=100, bottom=100, start=140, end=140)
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    for index, line in enumerate(text.splitlines()):
        if index:
            paragraph.add_run().add_break()
        run = paragraph.add_run(line)
        run.font.name = "Consolas"
        run._element.rPr.rFonts.set(qn("w:ascii"), "Consolas")
        run._element.rPr.rFonts.set(qn("w:hAnsi"), "Consolas")
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        run.font.size = Pt(8.5)
    set_table_geometry(table, [9360])


def add_paragraph(doc: Document, text: str, style: str | None = None, bold_label: str | None = None):
    paragraph = doc.add_paragraph(style=style)
    if bold_label and text.startswith(bold_label):
        label = paragraph.add_run(bold_label)
        set_run_font(label, 11, True)
        rest = paragraph.add_run(text[len(bold_label) :])
        set_run_font(rest, 11, False)
    else:
        run = paragraph.add_run(text)
        set_run_font(run, 11, False)
    return paragraph


def write_docx(results: list[dict[str, str | int | bool]]) -> Path:
    doc = Document()
    configure_document(doc)

    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(6)
    title_run = title.add_run("Scratch 算法题询问/回答测试文档")
    set_run_font(title_run, 22, True, "0B2545")

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(10)
    subtitle_run = subtitle.add_run(f"目标域：moran007.top / d/scratch    生成日期：{DOC_DATE}")
    set_run_font(subtitle_run, 10, False, "555555")

    add_paragraph(
        doc,
        "账号记录：管理员用户名 yang；学生用户名 ceshi1。密码不写入本地测试文档，请使用本次测试提供的口令。",
    )
    add_paragraph(
        doc,
        "测试约定：人工测试按 Scratch 的“询问并等待 / 答案”顺序输入；自动评测可把同一回答序列按换行拼成 input，最终只比较 output 的纯答案。",
    )

    doc.add_heading("本地参考运行结果", level=1)
    summary_rows = [["题号", "题目", "用例", "通过", "总分", "结果"]]
    for problem in PROBLEMS:
        problem_results = [row for row in results if row["pid"] == problem.pid]
        passed = sum(1 for row in problem_results if row["passed"])
        total_score = sum(int(row["score"]) for row in problem_results)
        summary_rows.append(
            [
                problem.pid,
                problem.title,
                str(len(problem_results)),
                f"{passed}/{len(problem_results)}",
                str(total_score),
                "PASS" if passed == len(problem_results) else "FAIL",
            ]
        )
    add_table(doc, summary_rows, [900, 2500, 900, 900, 900, 1200])
    add_paragraph(doc, f"本地参考求解器共校验 {len(results)} 个测试点，全部通过。")

    for problem in PROBLEMS:
        doc.add_heading(f"{problem.pid} {problem.title}", level=1)
        summary = [
            ["字段", "内容"],
            ["难度", problem.level],
            ["题意", problem.goal],
            ["询问顺序", "\n".join(problem.ask_flow)],
            ["输出要求", problem.output_rule],
            ["比较方式", problem.compare_mode],
            ["Scratch 提示", problem.scratch_tip],
        ]
        add_table(doc, summary, [1800, 7560])

        doc.add_heading("测试点", level=2)
        rows = [["用例", "回答序列", "期望输出", "分值", "结果", "说明"]]
        for case in problem.cases:
            actual = problem.solver(case.answers)
            rows.append(
                [
                    case.name,
                    "\n".join(case.answers),
                    case.expected,
                    str(case.score),
                    "PASS" if actual == case.expected else "FAIL",
                    "隐藏" if case.hidden else "公开",
                ]
            )
        add_table(doc, rows, [1700, 3000, 1500, 800, 900, 1460])

        doc.add_heading("快速录入", level=2)
        add_code_block(doc, "\n".join(quick_line(case) for case in problem.cases))

    path = OUT_DIR / f"{DOC_STEM}.docx"
    doc.save(path)
    return path


def main() -> None:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    results = validate_cases()
    md_path = write_markdown(results)
    docx_path = write_docx(results)
    json_path = write_json()
    quick_path = write_quick_entries()
    print(f"Generated Markdown: {md_path}")
    print(f"Generated DOCX: {docx_path}")
    print(f"Generated JSON: {json_path}")
    print(f"Generated quick entries: {quick_path}")
    print(f"Reference validation: {sum(1 for row in results if row['passed'])}/{len(results)} cases passed")


if __name__ == "__main__":
    main()
